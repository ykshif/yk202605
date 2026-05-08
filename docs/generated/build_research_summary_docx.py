from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = REPO_ROOT / "docs" / "RODM_300m_modular_connection_optimization_summary.docx"

PITCH_SUMMARY = (
    REPO_ROOT
    / "results"
    / "single_frequency_pitch_design_evaluation"
    / "pitch_single_frequency_evaluation_summary.csv"
)
DESIGN_SPACE_JSON = REPO_ROOT / "results" / "hinge_design_space" / "hinge_design_space_summary.json"
COMBINED_PARETO = (
    REPO_ROOT
    / "results"
    / "boundary18_combined_single_frequency"
    / "boundary18_combined_pareto_summary.csv"
)
COMBINED_FIGURE = (
    REPO_ROOT
    / "results"
    / "boundary18_combined_single_frequency"
    / "boundary18_combined_mean_heave_bending_pareto.png"
)
REFINED_FIGURE = (
    REPO_ROOT
    / "results"
    / "boundary18_refined_single_frequency"
    / "boundary18_doe_mean_heave_bending_pareto.png"
)
TRADEOFF_FIGURE = (
    REPO_ROOT
    / "results"
    / "single_frequency_pitch_design_evaluation"
    / "pitch_single_frequency_evaluation_summary.png"
)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def fmt(value: str | float, digits: int = 3) -> str:
    number = float(value)
    if number == 0:
        return "0"
    if abs(number) >= 1.0e4 or abs(number) < 1.0e-2:
        return f"{number:.{digits}e}"
    return f"{number:.{digits}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12.5, "385723"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_body_paragraph(document: Document, text: str, *, bold_prefix: str = ""):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("385723")
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_callout(document: Document, title: str, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    run.font.size = Pt(10.5)
    for line in lines:
        p = cell.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(2)
    document.add_paragraph()


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        cell.width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].width = Cm(widths[index])
    document.add_paragraph()


def add_figure(document: Document, path: Path, caption: str, width_inches: float = 6.2) -> None:
    if not path.exists():
        add_callout(document, "图像缺失", [str(path)])
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    caption_p = document.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.runs[0].font.size = Pt(9)
    caption_p.runs[0].font.color.rgb = RGBColor.from_string("666666")


def pareto_key_rows() -> list[dict[str, str]]:
    rows = read_csv(COMBINED_PARETO)
    pareto_rows = [row for row in rows if str(row.get("is_pareto", "")).lower() == "true"]
    names = [
        "uniform_high",
        "orient_x_1p00e09_y_3p00e08",
        "orient_x_1p00e09_y_1p80e08",
        "x_high_y_low",
        "center_stiff",
        "uniform_mid",
        "uniform_low",
    ]
    by_name = {row["design_label"]: row for row in pareto_rows}
    output = []
    for name in names:
        if name in by_name:
            output.append(by_name[name])
    return output


def build_document() -> None:
    document = Document()
    style_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("300m x 300m 大型浮体模块化连接优化研究阶段性总结")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("主问题：给定总尺度，优化模块数量、连接数量和连接刚度").bold = True
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("RODM 频域水弹性响应 · 连接件内力恢复 · 18 维连接刚度 DOE")

    add_callout(
        document,
        "本文档定位",
        [
            "这是一份面向论文写作的阶段性总结，不是最终论文全文。",
            "当前已完成的是固定 10 x 10 模块数量下的连接刚度优化子问题。",
            "后续论文应将模块数量、连接数量和连接刚度统一到同一个设计空间中。",
        ],
    )

    document.add_heading("1. 论文主问题与研究对象", level=1)
    add_body_paragraph(
        document,
        "论文主问题可以表述为：在 300m x 300m 总尺度约束下，确定模块划分方式、模块间连接数量和连接刚度分布，使大型浮体在给定海况下兼顾水弹性响应、连接件内力和结构可实现性。",
    )
    add_body_paragraph(
        document,
        "这个问题本质上不是单一刚度优化，而是一个“几何离散-连接拓扑-连接刚度”的耦合设计问题。模块数量决定单体尺度和界面数量，连接数量决定界面约束密度，连接刚度决定整体柔顺性与局部内力之间的权衡。",
    )
    add_table(
        document,
        ["设计层级", "变量示例", "对论文问题的意义"],
        [
            ["模块数量", "n x n，模块尺度 300/n m", "控制结构离散尺度、模块制造运输可行性和界面总数"],
            ["连接数量", "每条界面连接点数、连接线是否启用", "控制连接件数量、施工复杂度和局部传力路径"],
            ["连接刚度", "统一刚度、方向刚度、18 维边界刚度、180 维线刚度", "控制响应降低与连接件内力放大的折中关系"],
        ],
        [3.0, 4.5, 8.0],
    )

    document.add_heading("2. 当前代码与计算框架", level=1)
    add_body_paragraph(
        document,
        "当前工作已经从 notebook 原型逐步抽象出核心函数。src 目录承担结构构造、频域求解、连接件内力恢复和设计评价；notebooks 保留论文实验、图形展示和结果说明；scripts 用于可复现实验运行。",
    )
    add_table(
        document,
        ["模块", "核心职责", "当前状态"],
        [
            ["validation/complex_hinge_10x10.py", "构造并求解 10 x 10 铰接水弹性算例", "已用于所有单频试算"],
            ["strength/connector_recovery.py", "由 x_hat 恢复相对位移和连接件复内力", "已完成单元测试"],
            ["optimization/design_evaluator.py", "统一输出响应和连接件内力指标", "已支持 uniform pitch 与 boundary18"],
            ["optimization/hinge_design_space.py", "统计和映射铰接设计变量层级", "已支持 1/2/18/180/1260 维"],
            ["optimization/boundary18_doe.py", "生成 18 维 DOE 样本", "已完成 initial 与 refined 样本"],
        ],
        [4.6, 7.2, 3.8],
    )

    heading = document.add_heading("3. 10 x 10 模型中的连接数量", level=1)
    heading.paragraph_format.page_break_before = True
    add_body_paragraph(
        document,
        "对于当前 300m x 300m 总尺度、10 x 10 模块划分，每个模块尺寸为 30m x 30m。代码中每个模块边界含 7 个结构节点，因此一条模块间界面对应 7 个连接节点对。",
    )
    add_table(
        document,
        ["统计项", "数值", "说明"],
        [
            ["模块数量", "100", "10 x 10"],
            ["铰接线数量", "180", "x 方向 90 条，y 方向 90 条"],
            ["连接节点对数量", "1260", "180 条铰接线 x 每线 7 个节点对"],
            ["18 维边界变量", "18", "9 条纵向内部边界 + 9 条横向内部边界"],
            ["180 维线变量", "180", "每条铰接线一个刚度"],
            ["1260 维连接件变量", "1260", "每个节点对一个刚度，维度过高，暂不建议直接优化"],
        ],
        [4.0, 2.5, 8.8],
    )
    add_body_paragraph(
        document,
        "因此，18 维 boundary 参数化是当前阶段比较合适的折中：它保留了空间分布自由度，又避免直接进入 180 维或 1260 维高维优化。",
    )

    document.add_heading("4. 连接件内力恢复理论与实现", level=1)
    add_body_paragraph(
        document,
        "给定频域响应 x_hat，连接件相对位移和复内力按照以下关系恢复：delta_hat = G @ x_hat，force_hat = (K + i omega C) @ delta_hat。这里 G 是连接件相对位移算子，K 是连接刚度矩阵，C 是可选阻尼矩阵，omega 单位为 rad/s。",
    )
    add_body_paragraph(
        document,
        "对于优化目标，不直接取复力实部，而使用谐波合力包络 max_phi ||real(force_hat exp(i phi))||。这避免了相位选择导致的目标函数不一致，是后续连接件强度约束和多目标优化的基础。",
    )
    add_table(
        document,
        ["验证项", "验证结论"],
        [
            ["两节点三自由度弹簧", "k=1000 N/m，delta=0.01 m 时 force=10 N"],
            ["两端相同位移", "delta=0，force=0"],
            ["零刚度", "force=0"],
            ["组装一致性", "sum G.T @ force = Zc @ x"],
            ["复数响应与多频批量", "均已通过单元测试"],
        ],
        [5.2, 10.2],
    )

    document.add_heading("5. 固定单频下的已有数值结果", level=1)
    add_body_paragraph(
        document,
        "当前所有设计评价均固定在 omega = 0.5851 rad/s，波向为已有的 0 度水动力数据。频率扫描暂未开展，这是为了先把单频设计变量、连接件内力恢复和 Pareto 评价链条打通。",
    )
    pitch_rows = read_csv(PITCH_SUMMARY)
    add_table(
        document,
        ["pitch 刚度", "mean heave", "max heave", "max shear", "max bending"],
        [
            [
                row["pitch_stiffness_label"],
                fmt(row["mean_heave"]),
                fmt(row["max_heave"]),
                fmt(row["max_connector_shear_envelope"]),
                fmt(row["max_connector_bending_envelope"]),
            ]
            for row in pitch_rows
        ],
        [2.4, 2.8, 2.8, 3.2, 3.2],
    )
    add_figure(document, TRADEOFF_FIGURE, "图 1 固定单频下 uniform pitch 刚度的响应与连接件内力权衡", 6.1)

    document.add_heading("6. 18 维边界刚度 DOE 结果", level=1)
    add_body_paragraph(
        document,
        "在 18 维参数化中，每个变量控制一条完整内部边界。已完成 initial DOE 和 refined DOE 两轮样本。initial DOE 用于发现趋势，refined DOE 用于围绕 x 高刚度、y 低到中等刚度区域加密。",
    )
    key_rows = pareto_key_rows()
    add_table(
        document,
        ["设计", "样本集", "mean heave", "max bending", "max shear", "结论"],
        [
            [
                row["design_label"],
                row["sample_set"].replace("boundary18_", "").replace("_single_frequency", ""),
                fmt(row["mean_heave"]),
                fmt(row["max_connector_bending_envelope"]),
                fmt(row["max_connector_shear_envelope"]),
                {
                    "uniform_high": "响应最低但弯矩较高",
                    "orient_x_1p00e09_y_3p00e08": "低响应端有竞争力",
                    "orient_x_1p00e09_y_1p80e08": "较好的低响应折中",
                    "x_high_y_low": "接近低响应且弯矩略低",
                    "center_stiff": "中间权衡点",
                    "uniform_mid": "均衡基准点",
                    "uniform_low": "弯矩最低但响应较高",
                }.get(row["design_label"], ""),
            ]
            for row in key_rows
        ],
        [4.3, 2.8, 2.4, 2.8, 2.8, 3.8],
    )
    add_figure(document, COMBINED_FIGURE, "图 2 initial + refined 18 维 DOE 的合并 Pareto 结果", 6.2)
    add_figure(document, REFINED_FIGURE, "图 3 refined 18 维 DOE 在低响应区域的加密结果", 6.2)

    document.add_heading("7. 对主论文问题的启示", level=1)
    add_body_paragraph(
        document,
        "目前得到的最重要启示是：连接刚度不宜只用一个全局变量描述。x 方向边界高刚度、y 方向低到中等刚度的设计，在当前波向和频率下接近 uniform high 的低响应表现，同时连接件弯矩略有降低。",
    )
    add_body_paragraph(
        document,
        "这说明论文主问题应采用分层优化路线。第一层确定模块数量 n；第二层确定连接数量或连接密度；第三层在合理低维参数化下优化连接刚度。直接对 1260 个连接件刚度做优化不仅计算维度过高，而且不利于工程解释。",
    )
    add_table(
        document,
        ["阶段", "建议设计变量", "原因"],
        [
            ["阶段 A", "n 与 uniform/orientation 刚度", "建立模块数量与刚度尺度的总体趋势"],
            ["阶段 B", "18 维边界刚度", "保留空间分布，变量维度仍可控"],
            ["阶段 C", "局部 180 维细化", "仅在敏感边界附近展开，而非全局高维搜索"],
            ["阶段 D", "连接数量变量", "改变每条界面的连接点数，评估连接件数量与响应/内力的关系"],
        ],
        [2.2, 5.0, 8.2],
    )

    document.add_heading("8. 下一步论文工作计划", level=1)
    add_table(
        document,
        ["优先级", "任务", "预期输出"],
        [
            ["P1", "建立 2 维 x/y 刚度响应面", "验证方向刚度是否足以解释大部分收益"],
            ["P1", "引入连接件强度约束", "从 Pareto 点中筛出工程可行区域"],
            ["P2", "扩展模块数量 n 的试算", "比较 5x5、6x6、10x10 等模块划分"],
            ["P2", "连接数量参数化", "每条界面连接点数变化下的响应与内力曲线"],
            ["P3", "多频或鲁棒验证", "确认单频最优方案不只对一个频率有效"],
        ],
        [2.0, 6.2, 7.0],
    )

    document.add_heading("9. 阶段性结论", level=1)
    add_body_paragraph(
        document,
        "当前工作已经完成从 RODM 频域响应、连接件内力恢复、单频设计评价，到 18 维连接刚度 DOE 的完整链条。它为论文主问题奠定了基础：后续可以把模块数量和连接数量纳入同一套评价框架，而不需要推翻已有代码。",
    )
    add_body_paragraph(
        document,
        "在论文写作上，建议将当前 10 x 10 结果定位为“固定模块数量与连接数量下的连接刚度优化子问题”，并进一步扩展到“模块数量-连接数量-连接刚度协同优化”。这样主线清楚，也能解释为什么先从 18 维刚度参数化开始。",
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
